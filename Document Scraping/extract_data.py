import getopt
import pandas as pd
import json
import numpy as np
from docx.api import Document
import docx2txt
import sys
import os
from functools import reduce
from operator import mul


def reshape(lst, shape):
    if len(shape) == 1:
        return lst
    n = reduce(mul, shape[1:])
    return [reshape(lst[i*n:(i+1)*n], shape[1:]) for i in range(len(lst)//n)]


def get_personal_risk_insurance_policy(filename, name):
    try:
        personal_risk_insurance_policy = getData(filename, name)
        columns = ['Policy No.', 'Life & Amount', 'Type', 'Renewal Date', 'Issue Status', 'Stand Alone?',
                   'Buy Back?', 'Reinstatement?', 'Waiting Period', 'Benefit Period', 'Via Super', 'Annual Premium']
        sections = ['Life', 'TPD', 'Trauma',
                    'Income Protection', 'Business Expense']
        arr,section_order = fixData4(personal_risk_insurance_policy, columns, sections)
        owner = arr[0][1].split("\n")[0].strip()

        rev_section_order = {v:k for k,v in section_order.items()}
        # print(rev_section_order)

        # print(arr)
        # print(section_order)
        # print(owner)
        # for index, row in df.iterrows():
        #     if str(row['Life & Amount']).strip() != "":
        #         if len(str(row['Life & Amount']).split("\n")) >= 2 and row['Life & Amount'].split("\n")[1].strip().replace("$", "").replace(",", "").replace("p/m", "").strip().isnumeric():
        #             owner = row['Life & Amount'].split("\n")[0].strip()
        #             break

        out = []
        j = 0
        for i in range(len(arr)):
            if(i in rev_section_order):
                out.append({'section':rev_section_order[i],'data':[]})
                j+=1

            data = {columns[k]:arr[i][k] for k in range(len(columns))}
            out[j-1]['data'].append(data)
            
        # print(out)
        # out = processDf(df, sections, owner)
    except:
        return {}
    return out


def getData(filename, searchText):
    new_file_abs = filename
    document = Document(new_file_abs)
    check_second = False
    table = None
    if searchText == "Children / Dependants":
        for tb in document.tables:
            if len(tb.rows[0].cells) > 3:
                if tb.rows[0].cells[0].text == "Name" and tb.rows[0].cells[2].text == "Relationship" and tb.rows[0].cells[4].text == "Date of Birth":
                    table = tb
                    break
    elif searchText == "Personal Risk Insurance Cover":
        for index, tb in enumerate(document.tables):
            if len(tb.rows[0].cells) > 3:
                if tb.rows[0].cells[0].text.strip() == "Policy No." and tb.rows[0].cells[1].text.strip() == "Underwriter" and tb.rows[0].cells[2].text.strip() == "Policy Name":
                    table = tb
                    check_second = True
                    break
    elif searchText == "Personal Risk Insurance Cover 2":
        is_first = 0
        for index, tb in enumerate(document.tables):
            if len(tb.rows[0].cells) > 3:
                if tb.rows[0].cells[0].text.strip() == "Policy No." and tb.rows[0].cells[1].text.strip() == "Underwriter" and tb.rows[0].cells[2].text.strip() == "Policy Name":
                    if is_first == 0:
                        is_first = 1
                    else:
                        table = tb
                        check_second = True
                        break
    elif searchText == "Personal Risk Insurance Policy":
        for index, tb in enumerate(document.tables):
            if len(tb.rows[0].cells) > 3:
                if tb.rows[0].cells[0].text == "Policy No." and tb.rows[0].cells[1].text == "Life & Amount" and tb.rows[0].cells[2].text == "Type":
                    table = tb
                    break

    elif searchText == "Personal Risk Insurance Policy 2":
        is_first = 0
        for index, tb in enumerate(document.tables):
            if len(tb.rows[0].cells) > 3:
                if tb.rows[0].cells[0].text == "Policy No." and tb.rows[0].cells[1].text == "Life & Amount" and tb.rows[0].cells[2].text == "Type":
                    if is_first == 0:
                        is_first = 1
                    else:
                        table = tb
                        break
    else:
        for index, tb in enumerate(document.tables):
            if tb.rows[0].cells[0].text == searchText:
                table = tb
                break

    if check_second is True:
        df = [['' for i in range(len(table.columns))]
              for j in range(len(table.rows))]
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if cell.text:
                    df[i][j] = cell.text
        df = pd.DataFrame(df)
        df.columns = df.iloc[0]
        df = df[1:]
        return df

    if table is None:
        return None
    #table = document.tables[j]
    # return table

    data = []
    keys = None
    for i, row in enumerate(table.rows):
        text = [cell.text for cell in row.cells]
        if i == 0:
            keys_a = tuple(text)
            continue
        row_data = dict(zip(keys_a, text))
        data.append(row_data)

    df = pd.DataFrame(data)
    return df


def fixData4(df, columns, sections):
    print('inside fixdata4()')
    l = len(columns)
    all_data = []

    s_list = []
    s_order = {}
    f_s = False
    if len(df.columns) > l:
        f_s = True
        s_list.extend(df.columns[(len(df.columns) - l)*-1:])
        s_order[s_list[0]] = 0

    carry_over = None
    count = 0
    for index, row in df.iterrows():        
        r_v = list(row.values)
        # print(r_v,carry_over)
        if(f_s):
            f_s = False
            carry_over = r_v.pop()
            continue
        if carry_over in sections:
            f_s = True
            s_list.append(carry_over)
            carry_over = ''
            s_order[s_list[len(s_list)-1]] = count
            # print('after',r_v,carry_over)
        else:            
            r_v.insert(0,carry_over)
            carry_over = r_v.pop()
            r_v = r_v[0:l]

        count+=1
        all_data.append(r_v)

    # for index, p in enumerate(all_data):
    #     if p in sections:
    #         break
    all_data = all_data[0:]
   
    # df = pd.DataFrame(reshape(
    #     all_data, [int(len(all_data)/len(columns)), len(columns)]), columns=columns)
    # print(df)
    return all_data,s_order


def fixData5(df, columns):
    all_data = []

    if len(df.columns) > len(columns):
        all_data.extend(df.columns[(len(df.columns) - len(columns))*-1:])

    for index, row in df.iterrows():
        all_data.extend(list(row.values))

    df = pd.DataFrame(reshape(
        all_data, [int(len(all_data)/len(columns)), len(columns)]), columns=columns)
    return df


def processDf(df, sections, owner):
    columns = list(df.columns)
    c_section = ""
    out = {}
    c_data = []
    for index, row in df.iterrows():
        to_get = columns  # ["Policy No.","Via Super","Annual Premium"]
        to_replace = {"Policy No.": "policy_number",
                      "Via Super": "via_super", "Annual Premium": "annual_premium"}

        # extending to_replace to accomodate all the keys
        for x in to_get:
            if x in list(to_replace.keys()):
                pass
            else:
                to_replace[x] = x

        c_row = {to_replace[r]: row[r] for r in to_get}
        c_row.update({"owner": owner})

        if row[columns[0]] in sections:
            continue
        else:
            try:
                if len(c_row[to_replace[to_get[0]]].split("\n")) > 1 and c_row[to_replace[to_get[0]]].split("\n")[1].strip():
                    c_row[to_replace[to_get[0]]] = c_row[to_replace[to_get[0]]].split("\n")[
                        1].strip()

                if len(c_row[to_replace[to_get[2]]].split("\n")) > 1:
                    c_row[to_replace[to_get[2]]] = c_row[to_replace[to_get[2]]].split("\n")[
                        0].strip()
            except:
                pass
            if c_row is not None:
                c_row = {key: str(c_row[key]) for key in c_row.keys()}
                c_data.append(dict(c_row))

    return c_data


def getTable(table):
    data = []
    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    df = pd.DataFrame(data)
    return df


def docx_to_text(docx_file=""):
    # this function returns: error, extracted_text
    print("docx_file = " + docx_file)
    try:
        if not os.path.exists(docx_file):
            error_output = {}
            error_output['error'] = "File not found."
            return error_output, None

        filename = docx_file
        filename = filename.replace("\\", "/")
        text = docx2txt.process(filename)
        return None, text
    except Exception as e:
        print(f"error = {str(e)}")
        return str(e), None


def parse_text(filename, text=""):
    # this function returns: error, parsed_text
    all_details = {}
    personal_profile_dict = {}
    personal_details_dict = []
    # personal details
    try:
        document = Document(filename)
        names = list(getTable(document.tables[2]).columns)
        if "Title" in names and len(names) <= 3:
            indices = [0]
        else:
            indices = [0, 2]
        sep = "\n\n\n\n".join([names[i] for i in indices])+"\n\n"
        personal_details = []
        if "Title" in names and len(names) <= 3:
            for i in text.split("Personal Information Summary")[1].split("Children / Dependants")[0].strip().replace(sep, "", 1).strip().split("\n"*16):
                i = i.replace("\n"*4, "\t"*4)
                personal_details.extend(i.split("\n"*2))
        else:
            for i in text.split("Personal Information Summary")[1].split("Children / Dependants")[0].strip().replace(sep, "", 1).strip().replace("\n\n\n\n", "\t\t\t\t").split("\n\n"):
                personal_details.extend(i.split("\t\t\t\t\t\t\t\t"))
        personal_details = personal_details[:14]

        title_surname = personal_details[1].split("\t"*4)
        sex_marital_status = personal_details[7].split("\t"*4)
        dob_age_anb = personal_details[9].split("\t"*4)

        if len(indices) == 1:
            titles = [title_surname[0]]
            surnames = [title_surname[1]]
            given_names = [personal_details[3].split("\t"*4)[0]]
            preferred_name = [personal_details[3].split("\t"*4)[1]]
            sex = [sex_marital_status[0]]
            marital_status = [sex_marital_status[1]]
            dateofbirth = [dob_age_anb[0]]
        else:
            titles = [title_surname[0], title_surname[2]]
            surnames = [title_surname[1], title_surname[3]]
            given_names = [personal_details[3].split(
                "\t"*4)[0], personal_details[3].split("\t"*4)[2]]
            preferred_name = [personal_details[3].split(
                "\t"*4)[1], personal_details[3].split("\t"*4)[3]]
            sex = [sex_marital_status[0], sex_marital_status[2]]
            marital_status = [sex_marital_status[1], sex_marital_status[3]]
            dateofbirth = [dob_age_anb[0], dob_age_anb[3]]

        if len(surnames) == 1:
            personal_details_1 = {'Client': 'Client 1', 'title': titles[0], 'surname': surnames[0], 'given_name': given_names[0],
                                  'preferred_name': preferred_name[0], 'sex': sex[0], 'dateofbirth': dateofbirth[0], 'marital_status': marital_status[0]}
            personal_details_dict = [personal_details_1]
        elif len(surnames) == 2:
            num_clients = 2
            personal_details_1 = {'Client': 'Client 1',  'title': titles[0], 'surname': surnames[0], 'given_name': given_names[0],
                                  'preferred_name': preferred_name[0], 'sex': sex[0], 'dateofbirth': dateofbirth[0], 'marital_status': marital_status[0]}
            personal_details_2 = {'Client': 'Client 2', 'title': titles[1], 'surname': surnames[1], 'given_name': given_names[1],
                                  'preferred_name': preferred_name[1], 'sex': sex[1], 'dateofbirth': dateofbirth[1], 'marital_status': marital_status[1]}
            personal_details_dict = [personal_details_1, personal_details_2]
    except:
        print("Personal details not found in accepted format!")

    personal_risk_insurance = {}

    # personal_risk_insurance_policy
    personal_risk_insurance_policy_dict = []
    try:

        out = get_personal_risk_insurance_policy(
            filename, "Personal Risk Insurance Policy")
        personal_risk_insurance_policy_dict.extend(out)
        out2 = get_personal_risk_insurance_policy(
            filename, "Personal Risk Insurance Policy 2")
        if out2 != []:
            print("got            2")
            personal_risk_insurance_policy_dict.extend(out2)

    except:
        print("personal_risk_insurance_policy not found in accepted format!")
        pass
    personal_risk_insurance['personal_risk_insurance_details'] = personal_risk_insurance_policy_dict

    # personal_risk_insurance_cover
    personal_risk_insurance_cover_dict = []
    try:
        personal_risk_insurance_cover = getData(
            filename, "Personal Risk Insurance Cover")
        # print(personal_risk_insurance_cover,'\n')
        columns = ['Policy No', 'Underwriter', 'Policy Name', 'Life',
                   'TPD', 'Trauma', 'Income Protection', 'Business Expense']
        df = fixData5(personal_risk_insurance_cover, columns)
        # print(df,'\n')
        df = df[1:-1]
        owner = personal_risk_insurance_cover.columns[-1]
        df.columns = ['policy_number', 'underwriter', 'policy_name',
                      'Life', 'tpd', 'Trauma', 'income_protection', 'business_expense']
        df['owner'] = owner
        personal_risk_insurance_cover_dict.extend(
            json.loads(df.to_json(orient='records')))
        # print(personal_risk_insurance_cover_dict,'\n')
        try:
            personal_risk_insurance_cover = getData(
                filename, "Personal Risk Insurance Cover 2")
            df = fixData5(personal_risk_insurance_cover, columns)
            df = df[1:-1]

            owner = personal_risk_insurance_cover.columns[-1]
            df.columns = ['policy_number', 'underwriter', 'policy_name',
                          'Life', 'tpd', 'Trauma', 'income_protection', 'business_expense']
            df['owner'] = owner
            personal_risk_insurance_cover_dict.extend(
                json.loads(df.to_json(orient='records')))
        except:
            pass
    except:
        print("personal_risk_insurance_cover not found in accepted format!")
        pass
    personal_risk_insurance['personal_risk_insurance_cover'] = personal_risk_insurance_cover_dict

    all_details["personal_profile"] = personal_profile_dict
    all_details["personal_risk_insurance"] = personal_risk_insurance

    # this function returns: error, parsed_text
    return None, json.loads(json.dumps(all_details))


def main(argv):
    inputfile = ''
    outputfile = ''
    try:
        opts, args = getopt.getopt(argv, "hi:o:", ["ifile=", "ofile="])
    except getopt.GetoptError:
        print('test.py -i <inputfile> -o <outputfile>')
        sys.exit(2)
    for opt, arg in opts:
        if opt == '-h':
            print('test.py -i <inputfile> -o <outputfile>')
            sys.exit()
        elif opt in ("-i", "--ifile"):
            inputfile = arg
        elif opt in ("-o", "--ofile"):
            outputfile = arg
    print('Input file is ' + inputfile)
    print('Output file is ' + outputfile)

    _, extracted_text = docx_to_text("./" + inputfile)
    _, parsed_text = parse_text(filename=inputfile, text=extracted_text)
    with open(outputfile, "w") as outfile:
        json.dump(parsed_text, outfile, indent=4)


if __name__ == "__main__":
    main(sys.argv[1:])
